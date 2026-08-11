"""
Format Converter — 格式转换器
==============================
核心特性：
  - Mermaid → PNG/SVG（多种渲染方式）
  - Markdown → HTML（带样式）
  - Markdown → Word (.docx)
  - Markdown → PDF（可选）
  - 批量转换
  - 自动检测内容中的 Mermaid 图并渲染

渲染方式（按优先级）：
  1. mermaid-cli (mmdc) — 本地渲染，质量最好
  2. Kroki API — 在线渲染，无需安装
  3. mermaid.ink — 在线渲染，简单备用

用法：
    from scripts.format_converter import FormatConverter
    converter = FormatConverter()
    converter.mermaid_to_png("graph LR; A-->B", "output.png")
    converter.markdown_to_html("doc.md", "doc.html")
    converter.markdown_to_word("doc.md", "doc.docx")
"""
import logging
import os
import re
import subprocess
import tempfile
import time
import urllib.parse
import urllib.request
from pathlib import Path

logger = logging.getLogger(__name__)


class FormatConverter:
    """格式转换器"""

    def __init__(self, project_root: str = None):
        self._root = Path(project_root) if project_root else Path.cwd()
        self._mmdc_path = self._find_mmdc()
        self._pandoc_path = self._find_pandoc()

    # ─── 工具检测 ──────────────────────────────────

    def _find_mmdc(self) -> str | None:
        """查找 mermaid-cli"""
        for cmd in ["mmdc", "npx mmdc"]:
            try:
                result = subprocess.run(
                    cmd.split() + ["--version"],
                    capture_output=True, text=True, timeout=10,
                )
                if result.returncode == 0:
                    return cmd
            except Exception:
                pass
        return None

    def _find_pandoc(self) -> str | None:
        """查找 pandoc"""
        try:
            result = subprocess.run(
                ["pandoc", "--version"],
                capture_output=True, text=True, timeout=10,
            )
            if result.returncode == 0:
                return "pandoc"
        except Exception:
            pass
        return None

    # ─── Mermaid → PNG/SVG ────────────────────────

    def mermaid_to_png(self, mermaid_code: str, output_path: str,
                       width: int = 1200, theme: str = "default") -> bool:
        """Mermaid 代码 → PNG 图片

        Args:
            mermaid_code: Mermaid 语法代码
            output_path: 输出 PNG 路径
            width: 图片宽度
            theme: 主题 (default/dark/forest/neutral)

        Returns: True 成功, False 失败
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 方式 1: mermaid-cli (mmdc)
        if self._mmdc_path:
            return self._mmdc_render(mermaid_code, output_path, width, theme)

        # 方式 2: Kroki API
        if self._kroki_render(mermaid_code, output_path, "png"):
            return True

        # 方式 3: mermaid.ink
        if self._mermaid_ink_render(mermaid_code, output_path):
            return True

        logger.warning("Mermaid 渲染失败：无可用渲染方式（安装 mermaid-cli: npm i -g @mermaid-js/mermaid-cli）")
        return False

    def mermaid_to_svg(self, mermaid_code: str, output_path: str,
                       theme: str = "default") -> bool:
        """Mermaid 代码 → SVG 矢量图"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if self._mmdc_path:
            return self._mmdc_render(mermaid_code, output_path, 0, theme, fmt="svg")

        return bool(self._kroki_render(mermaid_code, output_path, "svg"))

    def _mmdc_render(self, code: str, output: Path, width: int,
                     theme: str, fmt: str = "png") -> bool:
        """使用 mermaid-cli 渲染"""
        try:
            with tempfile.NamedTemporaryFile(mode="w", suffix=".mmd",
                                             delete=False, encoding="utf-8") as f:
                f.write(code)
                mmd_file = f.name

            cmd = self._mmdc_path.split() + [
                "-i", mmd_file,
                "-o", str(output),
                "-t", theme,
                "-b", "transparent",
            ]
            if width > 0 and fmt == "png":
                cmd.extend(["-w", str(width)])

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            os.unlink(mmd_file)

            if result.returncode == 0:
                logger.info(f"mmdc 渲染成功: {output}")
                return True
            else:
                logger.warning(f"mmdc 渲染失败: {result.stderr}")
                return False
        except Exception as e:
            logger.warning(f"mmdc 渲染异常: {e}")
            return False

    def _kroki_render(self, code: str, output: Path, fmt: str) -> bool:
        """使用 Kroki API 渲染"""
        try:
            import base64
            import zlib

            # Kroki API: POST https://kroki.io/mermaid/{format}
            data = zlib.compress(code.encode("utf-8"), 9)
            encoded = base64.urlsafe_b64encode(data).decode("ascii")

            req = urllib.request.Request(
                f"https://kroki.io/mermaid/{fmt}/{encoded}",
                headers={"Accept": f"image/{fmt}"},
            )
            with urllib.request.urlopen(req, timeout=15) as resp:
                content = resp.read()

            with open(output, "wb") as f:
                f.write(content)
            logger.info(f"Kroki 渲染成功: {output}")
            return True
        except Exception as e:
            logger.debug(f"Kroki 渲染失败: {e}")
            return False

    def _mermaid_ink_render(self, code: str, output: Path) -> bool:
        """使用 mermaid.ink 渲染"""
        try:
            import base64
            encoded = base64.urlsafe_b64encode(code.encode("utf-8")).decode("ascii")
            url = f"https://mermaid.ink/img/{encoded}"

            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=15) as resp:
                content = resp.read()

            with open(output, "wb") as f:
                f.write(content)
            logger.info(f"mermaid.ink 渲染成功: {output}")
            return True
        except Exception as e:
            logger.debug(f"mermaid.ink 渲染失败: {e}")
            return False

    # ─── Markdown → HTML ──────────────────────────

    def markdown_to_html(self, md_path: str, html_path: str = None,
                         title: str = "文档", css: str = None) -> str:
        """Markdown → HTML

        Args:
            md_path: Markdown 文件路径
            html_path: 输出 HTML 路径（None 则返回 HTML 字符串）
            title: HTML 标题
            css: 自定义 CSS（None 则使用内置样式）

        Returns: HTML 内容字符串
        """
        md_path = Path(md_path)
        with open(md_path, encoding="utf-8") as f:
            md_content = f.read()

        html_body = self._markdown_to_html_body(md_content)

        if css is None:
            css = self._default_css()

        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>{css}</style>
</head>
<body>
    <div class="container">
{html_body}
    </div>
</body>
</html>"""

        if html_path:
            html_path = Path(html_path)
            html_path.parent.mkdir(parents=True, exist_ok=True)
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html)
            logger.info(f"HTML 已写入: {html_path}")

        return html

    def _markdown_to_html_body(self, md: str) -> str:
        """简单 Markdown → HTML 转换（不依赖外部库）"""
        lines = md.split("\n")
        html_lines = []
        in_code = False
        in_table = False
        in_list = False  # P1 修复：跟踪列表状态以正确包裹 <ul>

        for line in lines:
            # 代码块
            if line.strip().startswith("```"):
                if in_code:
                    html_lines.append("</code></pre>")
                    in_code = False
                else:
                    lang = line.strip()[3:].strip()
                    html_lines.append(f'<pre><code class="language-{lang}">')
                    in_code = True
                continue

            if in_code:
                html_lines.append(self._escape_html(line))
                continue

            # 标题
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                level = len(m.group(1))
                text = self._inline_md(m.group(2))
                html_lines.append(f"<h{level}>{text}</h{level}>")
                continue

            # 表格
            if "|" in line and line.strip().startswith("|"):
                cells = [c.strip() for c in line.split("|")[1:-1]]
                # P1 修复：空 cells 不应视为分隔行（all([]) == True 的陷阱）
                if cells and all(re.match(r"^[-:]+$", c) for c in cells):
                    continue  # 分隔行
                if not in_table:
                    html_lines.append("<table>")
                    in_table = True
                    html_lines.append("<thead><tr>")
                    for c in cells:
                        html_lines.append(f"<th>{self._inline_md(c)}</th>")
                    html_lines.append("</tr></thead><tbody>")
                else:
                    html_lines.append("<tr>")
                    for c in cells:
                        html_lines.append(f"<td>{self._inline_md(c)}</td>")
                    html_lines.append("</tr>")
                continue
            elif in_table:
                html_lines.append("</tbody></table>")
                in_table = False

            # 列表
            m = re.match(r"^[\s]*[-*+]\s+(.*)", line)
            if m:
                # P1 修复：列表项需用 <ul> 包裹，否则 HTML 不合法
                if not in_list:
                    html_lines.append("<ul>")
                    in_list = True
                html_lines.append(f"<li>{self._inline_md(m.group(1))}</li>")
                continue
            elif in_list:
                html_lines.append("</ul>")
                in_list = False

            # 引用
            if line.strip().startswith(">"):
                text = self._inline_md(line.strip()[1:].strip())
                html_lines.append(f"<blockquote>{text}</blockquote>")
                continue

            # 分隔线
            if re.match(r"^---+\s*$", line):
                html_lines.append("<hr>")
                continue

            # 普通段落
            if line.strip():
                html_lines.append(f"<p>{self._inline_md(line)}</p>")
            else:
                html_lines.append("")

        if in_table:
            html_lines.append("</tbody></table>")
        if in_list:
            html_lines.append("</ul>")
        if in_code:
            html_lines.append("</code></pre>")

        return "\n        ".join(html_lines)

    def _inline_md(self, text: str) -> str:
        """行内 Markdown 转换"""
        # 粗体
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        # 斜体
        text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
        # 行内代码
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        # 链接
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r'<a href="\2">\1</a>', text)
        return text

    def _escape_html(self, text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _default_css(self) -> str:
        return """
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
               line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }
        h1, h2, h3, h4, h5, h6 { margin-top: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }
        pre { background: #f6f8fa; padding: 16px; border-radius: 6px; overflow: auto; }
        code { background: #f6f8fa; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }
        pre code { background: none; padding: 0; }
        table { border-collapse: collapse; width: 100%; }
        th, td { border: 1px solid #ddd; padding: 8px 12px; }
        th { background: #f6f8fa; }
        blockquote { border-left: 4px solid #ddd; margin: 0; padding-left: 16px; color: #666; }
        hr { border: none; border-top: 2px solid #eee; }
        a { color: #0366d6; text-decoration: none; }
        """

    # ─── Markdown → Word ──────────────────────────

    def markdown_to_word(self, md_path: str, docx_path: str = None) -> str:
        """Markdown → Word (.docx)

        优先使用 pandoc（质量最好），回退到 python-docx。
        """
        md_path = Path(md_path)
        docx_path = md_path.with_suffix(".docx") if docx_path is None else Path(docx_path)

        docx_path.parent.mkdir(parents=True, exist_ok=True)

        # 方式 1: pandoc
        if self._pandoc_path:
            try:
                result = subprocess.run(
                    ["pandoc", str(md_path), "-o", str(docx_path),
                     "--from=markdown", "--to=docx"],
                    capture_output=True, text=True, timeout=60,
                )
                if result.returncode == 0:
                    logger.info(f"pandoc 转换成功: {docx_path}")
                    return str(docx_path)
                else:
                    logger.warning(f"pandoc 转换失败: {result.stderr}")
            except Exception as e:
                logger.warning(f"pandoc 转换异常: {e}")

        # 方式 2: python-docx
        try:
            from docx import Document
            doc = Document()
            with open(md_path, encoding="utf-8") as f:
                md_content = f.read()

            for line in md_content.split("\n"):
                m = re.match(r"^(#{1,6})\s+(.*)", line)
                if m:
                    level = min(len(m.group(1)), 9)
                    doc.add_heading(m.group(2), level=level)
                elif line.strip():
                    doc.add_paragraph(line.strip())

            doc.save(str(docx_path))
            logger.info(f"python-docx 转换成功: {docx_path}")
            return str(docx_path)
        except ImportError:
            logger.warning("python-docx 未安装（pip install python-docx）")
        except Exception as e:
            logger.warning(f"python-docx 转换失败: {e}")

        return ""

    # ─── 批量处理 ────────────────────────────────

    def render_mermaid_in_markdown(self, md_path: str, output_dir: str = None) -> str:
        """渲染 Markdown 中的所有 Mermaid 图为 PNG

        将 ```mermaid 代码块替换为图片引用。
        """
        md_path = Path(md_path)
        with open(md_path, encoding="utf-8") as f:
            content = f.read()

        output_dir = md_path.parent / "images" if output_dir is None else Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        def replace_mermaid(match):
            code = match.group(1)
            # P2 修复：同毫秒内多个图会重名，加计数器
            img_name = f"mermaid_{int(time.time()*1000)}_{replace_mermaid._counter}.png"
            replace_mermaid._counter += 1
            img_path = output_dir / img_name
            if self.mermaid_to_png(code, str(img_path)):
                return f"![{img_name}]({img_path})"
            else:
                return match.group(0)  # 渲染失败保留原样

        replace_mermaid._counter = 0

        content = re.sub(
            r"```mermaid\n(.*?)\n```",
            replace_mermaid,
            content,
            flags=re.DOTALL,
        )

        # 写回
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(content)

        logger.info(f"Mermaid 图渲染完成: {md_path}")
        return str(md_path)

    def status(self) -> dict:
        """获取转换器状态"""
        return {
            "mmdc": self._mmdc_path is not None,
            "pandoc": self._pandoc_path is not None,
            "capabilities": {
                "mermaid_to_png": True,
                "mermaid_to_svg": True,
                "markdown_to_html": True,
                "markdown_to_word": self._pandoc_path is not None,
            },
        }
